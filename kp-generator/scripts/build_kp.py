#!/usr/bin/env python3
"""Сборка коммерческого предложения из kp.json в PDF и DOCX.

Использование:
    python3 build_kp.py kp.json --pdf --docx
    python3 build_kp.py kp.json --out-dir ./out --html

Без флагов формата собирает оба файла.
PDF рендерится headless-браузером (Chrome/Chromium), DOCX — python-docx.
"""

import argparse
import html as htmlmod
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile

PLACEHOLDER_RE = re.compile(r"\[\[УТОЧНИТЬ:(.*?)\]\]|\[\[TODO:(.*?)\]\]", re.S)

LABELS = {
    "ru": {
        "kp": "Коммерческое предложение",
        "for": "Для",
        "valid": "Предложение действительно до",
        "situation": "Ситуация",
        "solution": "Что предлагаем",
        "scope": "Границы работ",
        "included": "Входит в работы",
        "excluded": "Не входит",
        "stages": "Этапы и сроки",
        "stage": "Этап",
        "duration": "Срок",
        "actions": "Что делаем",
        "client_input": "Что нужно от вас",
        "deliverable": "Результат этапа",
        "total": "Общий срок",
        "payment": "Оплата",
        "pricing": "Стоимость",
        "recommended": "Рекомендуем",
        "for_whom": "Кому подходит",
        "economics": "Экономика проекта",
        "proof": "Опыт и результаты",
        "task": "Задача",
        "result": "Результат",
        "faq": "Частые вопросы",
        "cta": "Следующий шаг",
        "cost": "Цена бездействия",
        "sources": "Источники данных",
    },
    "en": {
        "kp": "Proposal",
        "for": "Prepared for",
        "valid": "Valid until",
        "situation": "Where you are now",
        "solution": "What we propose",
        "scope": "Scope",
        "included": "Included",
        "excluded": "Not included",
        "stages": "Timeline",
        "stage": "Stage",
        "duration": "Duration",
        "actions": "What we do",
        "client_input": "What we need from you",
        "deliverable": "Deliverable",
        "total": "Total duration",
        "payment": "Payment",
        "pricing": "Investment",
        "recommended": "Recommended",
        "for_whom": "Best for",
        "economics": "Business case",
        "proof": "Track record",
        "task": "Challenge",
        "result": "Result",
        "faq": "FAQ",
        "cta": "Next step",
        "cost": "Cost of doing nothing",
        "sources": "Sources",
    },
}


# ---------------------------------------------------------------- helpers

def esc(text):
    if text is None:
        return ""
    return htmlmod.escape(str(text))


def mark(text):
    """Экранирует текст и подсвечивает плейсхолдеры [[УТОЧНИТЬ: ...]].

    Слово «УТОЧНИТЬ» остаётся видимым: без него в тексте появляются фразы
    вида «Здравствуйте, имя контакта», которые выглядят как готовый текст,
    и менеджер отправляет документ с дырой.
    """
    if text is None:
        return ""
    out = esc(text)
    out = re.sub(
        r"\[\[(?:УТОЧНИТЬ|TODO):(.*?)\]\]",
        lambda m: '<span class="todo">УТОЧНИТЬ: ' + m.group(1).strip() + "</span>",
        out,
        flags=re.S,
    )
    return out.replace("\n\n", "</p><p>").replace("\n", "<br>")


def find_placeholders(node, found):
    if isinstance(node, str):
        for m in PLACEHOLDER_RE.finditer(node):
            found.append((m.group(1) or m.group(2) or "").strip())
    elif isinstance(node, dict):
        for v in node.values():
            find_placeholders(v, found)
    elif isinstance(node, list):
        for v in node:
            find_placeholders(v, found)
    return found


def get(d, *path, default=None):
    cur = d
    for key in path:
        if not isinstance(cur, dict) or key not in cur:
            return default
        cur = cur[key]
    return cur if cur not in (None, "", [], {}) else default


def find_browser():
    candidates = [
        os.environ.get("CHROME_BIN"),
        "/opt/pw-browsers/chromium",
        shutil.which("google-chrome"),
        shutil.which("google-chrome-stable"),
        shutil.which("chromium"),
        shutil.which("chromium-browser"),
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
    ]
    for c in candidates:
        if c and os.path.exists(c):
            return c
    # playwright bundles
    for root in ("/opt/pw-browsers", os.path.expanduser("~/.cache/ms-playwright")):
        if os.path.isdir(root):
            for dirpath, _dirs, files in os.walk(root):
                for name in ("chrome", "headless_shell"):
                    if name in files:
                        p = os.path.join(dirpath, name)
                        if os.access(p, os.X_OK):
                            return p
    return None


# ---------------------------------------------------------------- HTML

CSS = """
@page { size: A4; margin: 16mm 15mm 16mm 15mm; }
* { box-sizing: border-box; }
body {
  font-family: 'Inter','Golos Text','Manrope','Segoe UI','Liberation Sans','DejaVu Sans',sans-serif;
  font-size: 10.5pt; line-height: 1.55; color: #14181F; margin: 0;
  -webkit-print-color-adjust: exact; print-color-adjust: exact;
}
h1,h2,h3 { margin: 0; font-weight: 700; letter-spacing: -0.01em; }
p { margin: 0 0 8px; }
ul { margin: 0 0 8px; padding-left: 18px; }
li { margin-bottom: 4px; }
.todo { background: #FFE8E8; color: #C0142E; border-bottom: 1.5px dashed #C0142E;
        padding: 0 3px; font-weight: 600; }

/* --- cover --- */
.cover { page-break-after: always; min-height: 246mm; display: flex;
         flex-direction: column; justify-content: space-between; }
.cover-top { display: flex; justify-content: space-between; align-items: flex-start;
             border-bottom: 3px solid var(--accent); padding-bottom: 10px; }
.brand { font-size: 15pt; font-weight: 800; color: var(--accent); letter-spacing: -0.02em; }
.brand small { display: block; font-size: 8.5pt; font-weight: 400; color: #6B7480;
               letter-spacing: 0; margin-top: 2px; }
.doc-meta { text-align: right; font-size: 8.5pt; color: #6B7480; line-height: 1.6; }
.cover-mid { padding: 34mm 0 0; }
.eyebrow { font-size: 9pt; text-transform: uppercase; letter-spacing: 0.12em;
           color: var(--accent); font-weight: 700; margin-bottom: 14px; }
.cover h1 { font-size: 26pt; line-height: 1.18; max-width: 92%; }
.cover .sub { font-size: 12pt; color: #47505C; margin-top: 16px; max-width: 88%; line-height: 1.45; }
.cover-for { margin-top: 30px; padding-top: 16px; border-top: 1px solid #E3E7EC;
             font-size: 10pt; color: #47505C; }
.cover-for b { color: #14181F; font-size: 12pt; display: block; margin-top: 3px; }
.cover-bot { font-size: 9pt; color: #6B7480; border-top: 1px solid #E3E7EC; padding-top: 10px; }
.cover-bot b { color: #14181F; }

/* --- sections --- */
/* Секции могут разрываться между страницами — иначе документ распухает
   пустотами. Неразрывными держим только смысловые единицы: карточку решения,
   таблицу, блок тарифов, кейс, вопрос FAQ, финальный призыв. */
.section { margin-bottom: 22px; }
.section-head { display: flex; align-items: baseline; gap: 10px; margin-bottom: 12px;
                border-bottom: 1.5px solid var(--accent); padding-bottom: 6px;
                page-break-after: avoid; }
.section-num { font-size: 9pt; font-weight: 700; color: var(--accent); }
.section-head h2 { font-size: 14pt; }
.lead { font-size: 11pt; color: #303842; margin-bottom: 12px; }

.point { margin-bottom: 10px; padding-left: 12px; border-left: 2px solid #E3E7EC; }
.point b { display: block; margin-bottom: 2px; }

.solution-block { margin-bottom: 14px; padding: 12px 14px; background: #F7F9FB;
                  border-radius: 6px; page-break-inside: avoid; }
.solution-block h3 { font-size: 11.5pt; margin-bottom: 5px; }
.result { margin-top: 7px; padding: 6px 10px; background: #fff; border-left: 3px solid var(--accent);
          font-size: 10pt; font-weight: 600; color: #14181F; border-radius: 0 4px 4px 0; }

.callout { padding: 11px 14px; background: #FFF7E6; border-left: 3px solid #E8A33D;
           border-radius: 0 4px 4px 0; margin: 12px 0; font-size: 10pt; }
.callout b { display: block; margin-bottom: 3px; }

table { width: 100%; border-collapse: collapse; font-size: 9.5pt; page-break-inside: avoid; }
th { text-align: left; background: var(--accent); color: #fff; padding: 8px 9px;
     font-weight: 600; font-size: 9pt; }
td { padding: 8px 9px; border-bottom: 1px solid #E3E7EC; vertical-align: top; }
tr:nth-child(even) td { background: #FAFBFC; }

.two-col { display: flex; gap: 14px; }
.two-col > div { flex: 1; padding: 12px 14px; border-radius: 6px; }
.col-in { background: #F2F8F3; border: 1px solid #D6E9D9; }
.col-ex { background: #FAFBFC; border: 1px solid #E3E7EC; }
.two-col h3 { font-size: 10pt; margin-bottom: 6px; }
.two-col ul { padding-left: 16px; font-size: 9.5pt; }

.tiers { display: flex; gap: 12px; align-items: stretch; page-break-inside: avoid; }
.tier { flex: 1; border: 1px solid #E3E7EC; border-radius: 8px; padding: 14px 13px;
        display: flex; flex-direction: column; }
.tier.rec { border: 2px solid var(--accent); box-shadow: 0 3px 14px rgba(0,0,0,.07); }
.tier-badge { display: inline-block; background: var(--accent); color: #fff; font-size: 7.5pt;
              font-weight: 700; text-transform: uppercase; letter-spacing: .08em;
              padding: 2px 7px; border-radius: 3px; margin-bottom: 7px; }
.tier h3 { font-size: 11.5pt; margin-bottom: 4px; }
.tier .whom { font-size: 8.5pt; color: #6B7480; margin-bottom: 9px; min-height: 26px; }
.tier .price { font-size: 15pt; word-break: break-word; font-weight: 800; color: var(--accent); letter-spacing: -0.02em; }
.tier .price-note { font-size: 8.5pt; color: #6B7480; margin-bottom: 9px; }
.tier ul { font-size: 9pt; padding-left: 15px; margin-top: 6px; }
.tier li { margin-bottom: 3px; }
.terms { margin-top: 12px; font-size: 9pt; color: #47505C; }
.terms li { margin-bottom: 2px; }

.econ { display: flex; flex-wrap: wrap; gap: 10px; margin-bottom: 10px; }
.econ-cell { flex: 1 1 30%; padding: 11px 13px; background: #F7F9FB; border-radius: 6px;
             border-top: 3px solid var(--accent); }
.econ-cell .label { font-size: 8.5pt; color: #6B7480; margin-bottom: 3px; }
.econ-cell .value { font-size: 14pt; font-weight: 800; letter-spacing: -0.02em; }

.case { padding: 12px 14px; border: 1px solid #E3E7EC; border-radius: 6px; margin-bottom: 10px;
        page-break-inside: avoid; }
.case h3 { font-size: 11pt; margin-bottom: 6px; }
.case .row { font-size: 9.5pt; margin-bottom: 3px; }
.case .row b { color: #6B7480; font-weight: 600; }
.quote { margin-top: 8px; padding-left: 11px; border-left: 3px solid var(--accent);
         font-style: italic; color: #303842; font-size: 9.5pt; }
.facts { display: flex; flex-wrap: wrap; gap: 8px; margin-top: 10px; }
.fact { background: #F7F9FB; border-radius: 5px; padding: 7px 11px; font-size: 9.5pt; }

.faq-item { margin-bottom: 10px; page-break-inside: avoid; }
.faq-item b { display: block; margin-bottom: 3px; }
.faq-item p { color: #303842; }

.cta { page-break-inside: avoid; margin-top: 8px; padding: 18px 20px; border-radius: 8px;
       background: var(--accent); color: #fff; }
.cta h2 { font-size: 15pt; margin-bottom: 8px; color: #fff; }
.cta p { color: rgba(255,255,255,.92); font-size: 10.5pt; }
.cta .action { margin-top: 10px; padding: 10px 14px; background: rgba(255,255,255,.16);
               border-radius: 5px; font-weight: 700; font-size: 11pt; }
.cta .todo { background: #fff; }
.signature { margin-top: 14px; padding-top: 12px; border-top: 1px solid #E3E7EC; font-size: 9.5pt; }
.signature b { font-size: 11pt; }
.signature .role { color: #6B7480; }
.footnote { margin-top: 16px; font-size: 8pt; color: #8A929C; line-height: 1.5;
            border-top: 1px solid #EDF0F3; padding-top: 8px; }
"""


def render_html(d):
    lang = get(d, "meta", "language", default="ru")
    L = LABELS.get(lang, LABELS["ru"])
    accent = get(d, "seller", "accent", default="#1B4DFF")
    parts = []
    n = [0]

    def head(key):
        n[0] += 1
        return (
            f'<div class="section-head"><span class="section-num">{n[0]:02d}</span>'
            f"<h2>{esc(L[key])}</h2></div>"
        )

    # ---------------- cover
    seller_name = get(d, "seller", "name", default="")
    tagline = get(d, "seller", "tagline", default="")
    meta_lines = []
    if get(d, "meta", "number"):
        meta_lines.append(esc(d["meta"]["number"]))
    if get(d, "meta", "date"):
        meta_lines.append(esc(d["meta"]["date"]))
    if get(d, "meta", "valid_until"):
        meta_lines.append(f'{esc(L["valid"])} {esc(d["meta"]["valid_until"])}')

    client_line = mark(get(d, "client", "company", default=""))
    contact = get(d, "client", "contact_name")
    if contact:
        role = get(d, "client", "contact_role")
        client_line += f'<br><span style="font-size:9.5pt;font-weight:400;color:#47505C">{mark(contact)}'
        client_line += f", {mark(role)}" if role else ""
        client_line += "</span>"

    mgr = get(d, "seller", "manager", default={}) or {}
    contacts = " · ".join(
        [x for x in [mark(mgr.get("name")), mark(mgr.get("phone")), mark(mgr.get("email")),
                     mark(get(d, "seller", "site"))] if x]
    )

    parts.append(f"""
<div class="cover">
  <div>
    <div class="cover-top">
      <div class="brand">{esc(seller_name)}{f"<small>{esc(tagline)}</small>" if tagline else ""}</div>
      <div class="doc-meta">{"<br>".join(meta_lines)}</div>
    </div>
    <div class="cover-mid">
      <div class="eyebrow">{esc(L["kp"])}</div>
      <h1>{mark(get(d, "cover", "headline", default=""))}</h1>
      {f'<div class="sub">{mark(d["cover"]["subheadline"])}</div>' if get(d, "cover", "subheadline") else ""}
      <div class="cover-for">{esc(L["for"])}<b>{client_line}</b></div>
    </div>
  </div>
  <div class="cover-bot">{contacts}</div>
</div>""")

    # ---------------- greeting
    if get(d, "greeting"):
        parts.append(f'<div class="section"><p class="lead">{mark(d["greeting"])}</p></div>')

    # ---------------- situation
    if get(d, "situation"):
        s = d["situation"]
        body = [head("situation")]
        if s.get("intro"):
            body.append(f'<p class="lead">{mark(s["intro"])}</p>')
        for p in s.get("points") or []:
            body.append(
                f'<div class="point"><b>{mark(p.get("title"))}</b>{mark(p.get("text"))}</div>'
            )
        if s.get("cost"):
            body.append(
                f'<div class="callout"><b>{esc(L["cost"])}</b>{mark(s["cost"])}</div>'
            )
        srcs = s.get("sources") or []
        if srcs:
            items = "; ".join(
                f'{esc(x.get("fact"))} — {esc(x.get("source"))}' for x in srcs
            )
            body.append(f'<div class="footnote">{esc(L["sources"])}: {items}</div>')
        parts.append(f'<div class="section">{"".join(body)}</div>')

    # ---------------- solution
    if get(d, "solution", "blocks"):
        body = [head("solution")]
        if get(d, "solution", "intro"):
            body.append(f'<p class="lead">{mark(d["solution"]["intro"])}</p>')
        for b in d["solution"]["blocks"]:
            block = f'<div class="solution-block"><h3>{mark(b.get("title"))}</h3><p>{mark(b.get("text"))}</p>'
            if b.get("result"):
                block += f'<div class="result">{mark(b["result"])}</div>'
            body.append(block + "</div>")
        parts.append(f'<div class="section">{"".join(body)}</div>')

    # ---------------- scope
    if get(d, "scope", "included") or get(d, "scope", "excluded"):
        inc = "".join(f"<li>{mark(x)}</li>" for x in (get(d, "scope", "included") or []))
        exc = "".join(f"<li>{mark(x)}</li>" for x in (get(d, "scope", "excluded") or []))
        cols = ""
        if inc:
            cols += f'<div class="col-in"><h3>{esc(L["included"])}</h3><ul>{inc}</ul></div>'
        if exc:
            cols += f'<div class="col-ex"><h3>{esc(L["excluded"])}</h3><ul>{exc}</ul></div>'
        parts.append(f'<div class="section">{head("scope")}<div class="two-col">{cols}</div></div>')

    # ---------------- stages
    if get(d, "stages", "items"):
        rows = ""
        for i, st in enumerate(d["stages"]["items"], 1):
            rows += (
                f"<tr><td><b>{esc(st.get('n', i))}. {mark(st.get('title'))}</b></td>"
                f"<td>{mark(st.get('duration'))}</td><td>{mark(st.get('actions'))}</td>"
                f"<td>{mark(st.get('client_input'))}</td><td>{mark(st.get('deliverable'))}</td></tr>"
            )
        tail = []
        if get(d, "stages", "total"):
            tail.append(f'<b>{esc(L["total"])}:</b> {mark(d["stages"]["total"])}')
        if get(d, "stages", "payment"):
            tail.append(f'<b>{esc(L["payment"])}:</b> {mark(d["stages"]["payment"])}')
        tail_html = f'<p style="margin-top:10px;font-size:10pt">{" &nbsp;·&nbsp; ".join(tail)}</p>' if tail else ""
        parts.append(f"""<div class="section">{head("stages")}
<table><thead><tr><th style="width:22%">{esc(L["stage"])}</th><th style="width:12%">{esc(L["duration"])}</th>
<th style="width:26%">{esc(L["actions"])}</th><th style="width:20%">{esc(L["client_input"])}</th>
<th style="width:20%">{esc(L["deliverable"])}</th></tr></thead><tbody>{rows}</tbody></table>{tail_html}</div>""")

    # ---------------- pricing
    if get(d, "pricing", "options"):
        body = [head("pricing")]
        if get(d, "pricing", "intro"):
            body.append(f'<p class="lead">{mark(d["pricing"]["intro"])}</p>')
        tiers = ""
        for o in d["pricing"]["options"]:
            rec = " rec" if o.get("recommended") else ""
            badge = f'<div class="tier-badge">{esc(L["recommended"])}</div>' if o.get("recommended") else ""
            items = "".join(f"<li>{mark(x)}</li>" for x in (o.get("includes") or []))
            tiers += f"""<div class="tier{rec}">{badge}<h3>{mark(o.get("name"))}</h3>
<div class="whom">{mark(o.get("for_whom"))}</div>
<div class="price">{mark(o.get("price"))}</div>
<div class="price-note">{mark(o.get("price_note"))}</div>
<ul>{items}</ul></div>"""
        body.append(f'<div class="tiers">{tiers}</div>')
        if get(d, "pricing", "terms"):
            terms = "".join(f"<li>{mark(x)}</li>" for x in d["pricing"]["terms"])
            body.append(f'<ul class="terms">{terms}</ul>')
        if get(d, "pricing", "note"):
            body.append(f'<p style="font-size:9.5pt;color:#47505C">{mark(d["pricing"]["note"])}</p>')
        parts.append(f'<div class="section">{"".join(body)}</div>')

    # ---------------- economics
    if get(d, "economics", "rows"):
        cells = "".join(
            f'<div class="econ-cell"><div class="label">{mark(r.get("label"))}</div>'
            f'<div class="value">{mark(r.get("value"))}</div></div>'
            for r in d["economics"]["rows"]
        )
        concl = (
            f'<div class="callout">{mark(d["economics"]["conclusion"])}</div>'
            if get(d, "economics", "conclusion") else ""
        )
        title = get(d, "economics", "title", default=L["economics"])
        n[0] += 1
        hd = (f'<div class="section-head"><span class="section-num">{n[0]:02d}</span>'
              f"<h2>{esc(title)}</h2></div>")
        parts.append(f'<div class="section">{hd}<div class="econ">{cells}</div>{concl}</div>')

    # ---------------- proof
    if get(d, "proof", "cases") or get(d, "proof", "facts"):
        body = [head("proof")]
        for c in get(d, "proof", "cases") or []:
            block = f'<div class="case"><h3>{mark(c.get("client"))}</h3>'
            if c.get("task"):
                block += f'<div class="row"><b>{esc(L["task"])}:</b> {mark(c["task"])}</div>'
            if c.get("result"):
                block += f'<div class="row"><b>{esc(L["result"])}:</b> {mark(c["result"])}</div>'
            if c.get("quote"):
                block += f'<div class="quote">{mark(c["quote"])}</div>'
            body.append(block + "</div>")
        if get(d, "proof", "facts"):
            facts = "".join(f'<div class="fact">{mark(f)}</div>' for f in d["proof"]["facts"])
            body.append(f'<div class="facts">{facts}</div>')
        parts.append(f'<div class="section">{"".join(body)}</div>')

    # ---------------- faq
    if get(d, "faq", "items"):
        items = "".join(
            f'<div class="faq-item"><b>{mark(q.get("q"))}</b><p>{mark(q.get("a"))}</p></div>'
            for q in d["faq"]["items"]
        )
        parts.append(f'<div class="section">{head("faq")}{items}</div>')

    # ---------------- cta
    if get(d, "cta"):
        c = d["cta"]
        block = f'<div class="cta"><h2>{esc(L["cta"])}</h2><p>{mark(c.get("text"))}</p>'
        if c.get("action"):
            block += f'<div class="action">{mark(c["action"])}</div>'
        if c.get("link"):
            block += f'<p style="margin-top:8px;font-size:9.5pt">{esc(c["link"])}</p>'
        block += "</div>"
        sig = ""
        if mgr:
            sig = (f'<div class="signature"><b>{esc(mgr.get("name"))}</b>'
                   f'<div class="role">{esc(mgr.get("role"))}, {esc(seller_name)}</div>'
                   f'<div>{" · ".join([x for x in [esc(mgr.get("phone")), esc(mgr.get("email"))] if x])}</div></div>')
        parts.append(f'<div class="section">{block}{sig}</div>')

    if get(d, "footer_note") or get(d, "seller", "legal_name"):
        note = " ".join(filter(None, [get(d, "seller", "legal_name", default=""),
                                      get(d, "footer_note", default="")]))
        parts.append(f'<div class="footnote">{mark(note)}</div>')

    title = f'{L["kp"]} — {get(d, "client", "company", default="")}'
    return f"""<!DOCTYPE html><html lang="{lang}"><head><meta charset="utf-8">
<title>{esc(title)}</title>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800&display=swap" rel="stylesheet">
<style>:root {{ --accent: {accent}; }}
{CSS}</style></head><body>{"".join(parts)}</body></html>"""


# ---------------------------------------------------------------- PDF

def build_pdf(html_path, pdf_path):
    browser = find_browser()
    if not browser:
        raise RuntimeError(
            "Не найден Chrome/Chromium для рендера PDF. Установи браузер или задай CHROME_BIN. "
            "Как запасной вариант можно отдать HTML и DOCX."
        )
    profile = tempfile.mkdtemp(prefix="kpchrome-")
    cmd = [
        browser, "--headless", "--disable-gpu", "--no-sandbox", "--no-pdf-header-footer",
        "--run-all-compositor-stages-before-draw", "--virtual-time-budget=6000",
        f"--user-data-dir={profile}", f"--print-to-pdf={pdf_path}",
        "file://" + os.path.abspath(html_path),
    ]
    res = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
    shutil.rmtree(profile, ignore_errors=True)
    if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) < 1000:
        raise RuntimeError(f"PDF не собрался.\n{res.stdout}\n{res.stderr}")
    return pdf_path


# ---------------------------------------------------------------- DOCX

def build_docx(d, docx_path):
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Cm

    lang = get(d, "meta", "language", default="ru")
    L = LABELS.get(lang, LABELS["ru"])
    accent_hex = (get(d, "seller", "accent", default="#1B4DFF") or "#1B4DFF").lstrip("#")
    accent = RGBColor.from_string(accent_hex.upper())
    red = RGBColor.from_string("C0142E")

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2)
        s.left_margin = s.right_margin = Cm(2)

    def para(text, size=10.5, bold=False, color=None, space_after=6, align=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        if align:
            p.alignment = align
        if text is None:
            return p
        for chunk in re.split(r"(\[\[(?:УТОЧНИТЬ|TODO):.*?\]\])", str(text), flags=re.S):
            if not chunk:
                continue
            m = re.match(r"\[\[(?:УТОЧНИТЬ|TODO):(.*?)\]\]", chunk, flags=re.S)
            r = p.add_run(f"[УТОЧНИТЬ: {m.group(1).strip()}]" if m else chunk)
            r.font.size = Pt(size)
            r.bold = bold or bool(m)
            r.italic = italic
            r.font.color.rgb = red if m else (color or RGBColor.from_string("14181F"))
        return p

    def heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(str(text).upper())
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = accent

    def bullets(items):
        for x in items:
            p = para("• " + str(x), size=10, space_after=2)
            p.paragraph_format.left_indent = Cm(0.5)

    # cover
    para(get(d, "seller", "name", default=""), size=16, bold=True, color=accent, space_after=2)
    if get(d, "seller", "tagline"):
        para(d["seller"]["tagline"], size=9, color=RGBColor.from_string("6B7480"))
    meta_bits = [get(d, "meta", "number", default=""), get(d, "meta", "date", default="")]
    if get(d, "meta", "valid_until"):
        meta_bits.append(f'{L["valid"]} {d["meta"]["valid_until"]}')
    para(" · ".join([x for x in meta_bits if x]), size=9, color=RGBColor.from_string("6B7480"), space_after=14)

    para(L["kp"].upper(), size=9, bold=True, color=accent, space_after=6)
    para(get(d, "cover", "headline", default=""), size=20, bold=True, space_after=8)
    if get(d, "cover", "subheadline"):
        para(d["cover"]["subheadline"], size=12, color=RGBColor.from_string("47505C"))
    who = get(d, "client", "company", default="")
    if get(d, "client", "contact_name"):
        who += f' · {d["client"]["contact_name"]}'
        if get(d, "client", "contact_role"):
            who += f', {d["client"]["contact_role"]}'
    para(f'{L["for"]}: {who}', size=11, bold=True, space_after=14)

    if get(d, "greeting"):
        para(d["greeting"], size=11)

    if get(d, "situation"):
        heading(L["situation"])
        if get(d, "situation", "intro"):
            para(d["situation"]["intro"])
        for p in get(d, "situation", "points") or []:
            para(p.get("title"), bold=True, space_after=1)
            para(p.get("text"), space_after=8)
        if get(d, "situation", "cost"):
            para(f'{L["cost"]}: {d["situation"]["cost"]}', bold=True)

    if get(d, "solution", "blocks"):
        heading(L["solution"])
        if get(d, "solution", "intro"):
            para(d["solution"]["intro"])
        for b in d["solution"]["blocks"]:
            para(b.get("title"), size=11.5, bold=True, space_after=2)
            para(b.get("text"), space_after=2)
            if b.get("result"):
                para(f'→ {b["result"]}', bold=True, color=accent, space_after=10)

    if get(d, "scope", "included") or get(d, "scope", "excluded"):
        heading(L["scope"])
        if get(d, "scope", "included"):
            para(L["included"], bold=True, space_after=2)
            bullets(d["scope"]["included"])
        if get(d, "scope", "excluded"):
            para(L["excluded"], bold=True, space_after=2)
            bullets(d["scope"]["excluded"])

    if get(d, "stages", "items"):
        heading(L["stages"])
        cols = [L["stage"], L["duration"], L["actions"], L["client_input"], L["deliverable"]]
        t = doc.add_table(rows=1, cols=len(cols))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, c in enumerate(cols):
            cell = t.rows[0].cells[i]
            cell.text = ""
            r = cell.paragraphs[0].add_run(c)
            r.bold = True
            r.font.size = Pt(9)
        for i, s in enumerate(d["stages"]["items"], 1):
            cells = t.add_row().cells
            vals = [f'{s.get("n", i)}. {s.get("title","")}', s.get("duration", ""),
                    s.get("actions", ""), s.get("client_input", ""), s.get("deliverable", "")]
            for j, v in enumerate(vals):
                cells[j].text = ""
                r = cells[j].paragraphs[0].add_run(
                    re.sub(r"\[\[(?:УТОЧНИТЬ|TODO):(.*?)\]\]", r"[УТОЧНИТЬ: \1]", str(v), flags=re.S))
                r.font.size = Pt(9)
                if "[УТОЧНИТЬ:" in r.text:
                    r.bold = True
                    r.font.color.rgb = red
        doc.add_paragraph()
        for key in ("total", "payment"):
            if get(d, "stages", key):
                para(f'{L[key]}: {d["stages"][key]}', bold=True, space_after=2)

    if get(d, "pricing", "options"):
        heading(L["pricing"])
        if get(d, "pricing", "intro"):
            para(d["pricing"]["intro"])
        opts = d["pricing"]["options"]
        t = doc.add_table(rows=1, cols=len(opts))
        t.style = "Light Grid Accent 1"
        for i, o in enumerate(opts):
            cell = t.rows[0].cells[i]
            cell.text = ""
            p0 = cell.paragraphs[0]
            r = p0.add_run(o.get("name", "") + (f' — {L["recommended"]}' if o.get("recommended") else ""))
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = accent
            pp = cell.add_paragraph()
            rp = pp.add_run(str(o.get("price", "")))
            rp.bold = True
            rp.font.size = Pt(13)
            if o.get("price_note"):
                pn = cell.add_paragraph().add_run(str(o["price_note"]))
                pn.font.size = Pt(8.5)
                pn.font.color.rgb = RGBColor.from_string("6B7480")
            if o.get("for_whom"):
                pw = cell.add_paragraph().add_run(str(o["for_whom"]))
                pw.font.size = Pt(9)
                pw.italic = True
            for inc in o.get("includes") or []:
                ri = cell.add_paragraph().add_run("• " + str(inc))
                ri.font.size = Pt(9)
        doc.add_paragraph()
        if get(d, "pricing", "terms"):
            bullets(d["pricing"]["terms"])
        if get(d, "pricing", "note"):
            para(d["pricing"]["note"], size=9.5)

    if get(d, "economics", "rows"):
        heading(get(d, "economics", "title", default=L["economics"]))
        for r_ in d["economics"]["rows"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            a = p.add_run(f'{r_.get("label","")}: ')
            a.font.size = Pt(10)
            b = p.add_run(str(r_.get("value", "")))
            b.bold = True
            b.font.size = Pt(11)
            b.font.color.rgb = accent
        if get(d, "economics", "conclusion"):
            para(d["economics"]["conclusion"], bold=True)

    if get(d, "proof", "cases") or get(d, "proof", "facts"):
        heading(L["proof"])
        for c in get(d, "proof", "cases") or []:
            para(c.get("client"), size=11, bold=True, space_after=2)
            if c.get("task"):
                para(f'{L["task"]}: {c["task"]}', space_after=2)
            if c.get("result"):
                para(f'{L["result"]}: {c["result"]}', space_after=2)
            if c.get("quote"):
                para(f'«{c["quote"]}»', italic=True, space_after=8)
        if get(d, "proof", "facts"):
            bullets(d["proof"]["facts"])

    if get(d, "faq", "items"):
        heading(L["faq"])
        for q in d["faq"]["items"]:
            para(q.get("q"), bold=True, space_after=1)
            para(q.get("a"), space_after=8)

    if get(d, "cta"):
        heading(L["cta"])
        para(get(d, "cta", "text", default=""), size=11)
        if get(d, "cta", "action"):
            para(get(d, "cta", "action"), size=12, bold=True, color=accent)
        if get(d, "cta", "link"):
            para(d["cta"]["link"], size=9.5)
    mgr = get(d, "seller", "manager", default={}) or {}
    if mgr:
        para("", space_after=4)
        para(mgr.get("name", ""), size=11, bold=True, space_after=1)
        para(f'{mgr.get("role","")}, {get(d, "seller", "name", default="")}', size=9.5,
             color=RGBColor.from_string("6B7480"), space_after=1)
        para(" · ".join([x for x in [mgr.get("phone"), mgr.get("email"),
                                     get(d, "seller", "site")] if x]), size=9.5)
    if get(d, "footer_note") or get(d, "seller", "legal_name"):
        para(" ".join(filter(None, [get(d, "seller", "legal_name", default=""),
                                    get(d, "footer_note", default="")])),
             size=8, color=RGBColor.from_string("8A929C"))

    doc.save(docx_path)
    return docx_path


# ---------------------------------------------------------------- main

def main():
    ap = argparse.ArgumentParser(description="Сборка КП из kp.json в PDF и DOCX")
    ap.add_argument("json_file")
    ap.add_argument("--out-dir", default=".")
    ap.add_argument("--pdf", action="store_true")
    ap.add_argument("--docx", action="store_true")
    ap.add_argument("--html", action="store_true", help="сохранить промежуточный HTML")
    ap.add_argument("--name", help="базовое имя выходных файлов")
    args = ap.parse_args()

    with open(args.json_file, encoding="utf-8") as f:
        try:
            data = json.load(f)
        except json.JSONDecodeError as e:
            sys.exit(f"Невалидный JSON в {args.json_file}: {e}")

    for req in ("seller", "client", "cover"):
        if req not in data:
            sys.exit(f"В kp.json нет обязательной секции '{req}'. Схема — assets/kp-schema.md")

    os.makedirs(args.out_dir, exist_ok=True)
    base = args.name or "КП_" + re.sub(
        r"[^\w\-]+", "_", str(get(data, "client", "company", default="client"))
    ).strip("_")
    stem = os.path.join(args.out_dir, base)

    want_pdf = args.pdf or not (args.pdf or args.docx)
    want_docx = args.docx or not (args.pdf or args.docx)

    html_path = stem + ".html" if args.html else os.path.join(
        tempfile.mkdtemp(prefix="kp-"), "kp.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(render_html(data))

    made = []
    if args.html:
        made.append(html_path)
    if want_pdf:
        try:
            made.append(build_pdf(html_path, stem + ".pdf"))
        except Exception as e:
            print(f"! PDF не собран: {e}", file=sys.stderr)
            if not args.html:
                fallback = stem + ".html"
                shutil.copy(html_path, fallback)
                made.append(fallback)
                print(f"  HTML сохранён рядом: {fallback}", file=sys.stderr)
    if want_docx:
        try:
            made.append(build_docx(data, stem + ".docx"))
        except ImportError:
            print("! DOCX не собран: нет python-docx (pip install python-docx --break-system-packages)",
                  file=sys.stderr)

    print("Готово:")
    for m in made:
        print(f"  {m}  ({os.path.getsize(m) // 1024} КБ)")

    todos = find_placeholders(data, [])
    if todos:
        print(f"\nНезаполненных мест: {len(todos)} — их видно красным в документе:")
        for t in todos:
            print(f"  • {t}")
    else:
        print("\nНезаполненных плейсхолдеров нет.")


if __name__ == "__main__":
    main()
